from io import BytesIO
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from app.database import SessionLocal
from app.main import app
from app.models import DocumentRecord
from app.tasks import process_document


def docx_bytes() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_heading("季度经营报告", level=1)
    document.add_paragraph("本季度销售收入增长百分之二十。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "指标", "结果"
    table.cell(1, 0).text, table.cell(1, 1).text = "收入", "增长20%"
    document.save(stream)
    return stream.getvalue()


def xlsx_bytes() -> bytes:
    stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "销售数据"
    sheet.append(["月份", "销售额"])
    sheet.append(["一月", 100])
    workbook.save(stream)
    return stream.getvalue()


def test_batch_upload_query_and_history(monkeypatch):
    monkeypatch.setattr("app.main.process_document.delay", lambda record_id: type("Task", (), {"id": f"task-{record_id}"})())
    with TestClient(app) as client:
        response = client.post("/api/v1/documents", files=[
            ("files", ("report.docx", docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ("files", ("sales.xlsx", xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
        ])
        assert response.status_code == 202
        uploaded = response.json()
        assert uploaded["total"] == 2

        for record in uploaded["records"]:
            process_document.run(record["id"])
            detail = client.get(f"/api/v1/documents/{record['id']}")
            assert detail.status_code == 200
            assert detail.json()["status"] == "SUCCESS"
            assert detail.json()["summary"]["conclusion"] == "部署链路验证成功。"

        batch = client.get(f"/api/v1/batches/{uploaded['batch_id']}")
        assert batch.status_code == 200
        assert batch.json()["success_count"] == 2

        history = client.get("/api/v1/documents?page=1&page_size=10&status=SUCCESS")
        assert history.status_code == 200
        assert history.json()["total"] == 2

        with SessionLocal() as db:
            records = list(db.query(DocumentRecord).all())
            assert all(record.temp_path is None for record in records)
            assert all(record.extracted_text is None for record in records)


def test_rejects_unsupported_file(monkeypatch):
    monkeypatch.setattr("app.main.process_document.delay", lambda _: None)
    with TestClient(app) as client:
        response = client.post("/api/v1/documents", files={"files": ("bad.txt", b"bad", "text/plain")})
        assert response.status_code == 415
